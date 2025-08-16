import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def convert_to_list_of_dict(data):
    pass

def download_as_word(mcq_data, output_file = f"MCQ_Questions_{current_time}.docx"):
    # Create document
    doc = Document()
    # Add title
    title = doc.add_heading('Multiple Choice Questions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format each question
    for qnum, data in mcq_data.items():
        # Add question
        para = doc.add_paragraph()
        para.add_run(f"\nQ{qnum}. ").bold = True
        para.add_run(f"{data['mcq']}\n")
        
        # Add topic and difficulty
        topic_diff = doc.add_paragraph()
        topic_diff.add_run(f"Topic: {data['Topic']} | ").italic = True
        topic_diff.add_run(f"Difficulty: {data['Difficulty']}\n").italic = True
        
        # Add options
        for opt, choice in data['options'].items():
            doc.add_paragraph(f"{opt}) {choice}")
        
        # Add answer and explanation
        ans_para = doc.add_paragraph()
        ans_para.add_run("\nAnswer: ").bold = True
        ans_para.add_run(f"{data['Answer']}")
        
        exp_para = doc.add_paragraph()
        exp_para.add_run("Explanation: ").bold = True
        exp_para.add_run(f"{data['Explanation']}\n")
        
        # Add separator
        doc.add_paragraph("_" * 50)
    
    # Save document
    doc.save(os.path.join(ROOT_DIR, 'outputs', output_file))
